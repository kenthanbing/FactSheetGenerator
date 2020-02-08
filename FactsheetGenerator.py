import docx, os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import json

doc = Document('FactSheetGenerator' + os.sep + 'letterhead.docx')

# 定义写入标题的函数
def set_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return

# 添加图片的函数
def add_pictures(pic_paragragh):
    pic_run = pic_paragragh.add_run()
    for i in range(1,4):
        pic_run.add_picture('FactSheetGenerator' + os.sep + f'{str(i)}.jpg', width=docx.shared.Cm(5.4),height=docx.shared.Cm(3.3))
        pic_run.add_text(' ')

# 定义写入段落的函数
def write_paragraph(paragraphs):
    for paragraph in paragraphs:
        paragraph = doc.add_paragraph(paragraph)
        paragraph.paragraph_format.first_line_indent = Inches(0.32)
    return

# 获取json文件数据
with open('FactSheetGenerator' + os.sep + 'factsheet_data.json', 'rb') as f:
    factsheet_data = json.load(f)

# 设置正文全局字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 写标题
doc._body.clear_content()
run = doc.add_heading(level=0).add_run(factsheet_data['title'])
set_font(run)

# 写展会时间，国家地区，展会地点，举办届数
runs = [
    doc.add_heading(level=3).add_run("展会时间：" + factsheet_data['date']),
    doc.add_heading(level=3).add_run("国家地区：" + factsheet_data['country']),
    doc.add_heading(level=3).add_run("展会地点：" + factsheet_data['venue']),
    doc.add_heading(level=3).add_run("举办届数：" + factsheet_data['edition'])
]
for run in runs:
    set_font(run)

# 添加图片
pic_paragraph = doc.add_paragraph()
add_pictures(pic_paragraph)

# 写展会简介
run = doc.add_heading(level=3).add_run("展会简介：")
set_font(run)
paragraphs = factsheet_data['desc']
write_paragraph(paragraphs)

# 写往届回顾
run = doc.add_heading(level=3).add_run("往届回顾：")
set_font(run)
paragraphs = factsheet_data['report']
write_paragraph(paragraphs)

# 写市场分析
run = doc.add_heading(level=3).add_run("市场分析：")
set_font(run)
paragraphs = factsheet_data['market']
write_paragraph(paragraphs)

# 写参展范围
run = doc.add_heading(level=3).add_run("参展范围：")
set_font(run)
paragraphs = factsheet_data['exhibits']
write_paragraph(paragraphs)

# 写参展费用
run = doc.add_heading(level=3).add_run("参展费用：")
set_font(run)
paragraphs = factsheet_data['price']
write_paragraph(paragraphs)

# 写联系方式
run = doc.add_heading(level=3).add_run("联系方式：")
set_font(run)
paragraphs = factsheet_data['contact']
write_paragraph(paragraphs)

doc.save('FactSheetGenerator' + os.sep + f'{factsheet_data["title"]}.docx')
